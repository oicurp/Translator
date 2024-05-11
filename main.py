# -*- coding: utf-8 -*-
"""
Created on Fri Aug 18 09:22:23 2023

@author: Danny Huang

Notes: 
    1）为了使用脚本文件，用户必须预先在百度翻译网站注册一个开发者账户，然后开通“通用领域”和“垂直领域”服务。
    2）本脚本运行的环境是Anaconda包装的Python，使用脚本前，需要用conda安装一个docx模块，用来操作word文件
    
Development needed:
    1)Word文件的页眉和页脚还不能翻译，需要增加代码，用于遍历document.header和document.footer对象里的paragraph和table。
    2）Word文件在保存的时候，会丢失文本格式，这个需要比较多的代码解决。思路：提取段落的格式，保存的时候，把格式应用在文本上。

version 20240423: replaced "||" with "\n" to insert line break for english text.
"""
import os

from docx import Document

import translate as t # 通用领域的翻译模块，根据需要选择，更专业的，比如QC选用下面这个模块
#import fieldtranslate as t # 医疗领域的翻译模块，建议分析方法选择本模块。


source_folder = r"D:\Automation\Scripts\Translate\1_source"
output_folder = r"D:\Automation\Scripts\Translate\2_output"
file_path = ""
file_out_path = ""
file_name_english = ""
temp_para = ""
cell_text = []

def contain_cc(feed_text):
    for char in feed_text:
        if ord(char) > 10000:
            return True
        else:
            return False
        
    

# Loop through source files

for file in os.listdir(source_folder):
    file_path = source_folder + "\\" + file
    print(file_path)
    # open the first file for translation
    document = Document(file_path)
    
    # deal with text inside paragraphs
    for p in document.paragraphs:
        if p.text == temp_para:
            continue
        if p.text.isdigit():    # make sure numbers are not translated into english
            continue
        elif p.text.replace('.', '', 1).isdigit():
            continue
       # elif not contain_cc(p.text):
        #    continue
        else:
            temp_para = p.text
            p.text = p.text + "\n" + t.translate_text(p.text) # code can be changed to save english text only
           
        
    # deal with text inside tables
    for table in document.tables:
        for r in table.rows:
            cell_text=""
            for c in r.cells:
                if c.paragraphs == cell_text:
                    continue
                temp_para = ""
                for p in c.paragraphs:
                    if "||" in p.text:
                        continue
                    if p.text == temp_para:
                        continue
                    if p.text == temp_para: # this is to avoid translating merged cells, which are regarded as seprate cells in codes.
                        continue
                    if p.text.isdigit(): # need more control here, e.g., digit, float, or alphabet, do not translate.
                        continue
                    elif p.text.replace('.', '', 1).isdigit():
                        continue
                    elif not contain_cc(p.text):
                        continue
                    else:
                        p.text = p.text + "\n" + t.translate_text(p.text) # code can be changed to save english text only
                    temp_para = p.text
                cell_text = c.paragraphs
    
    #file_name_english = t.translate_text(file)
    #file_name_english = file_name_english.replace('. docx','.docx') # fixing file extension
    file_out_path = os.path.join(output_folder, file) # ensure file name is also translated
    document.save(file_out_path)


print("---------------------------------------")
print("Translation successfully completed.")